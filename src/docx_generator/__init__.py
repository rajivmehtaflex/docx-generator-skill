"""
DocX Generator - Advanced HTML to DOCX Converter
Browser-based Word document generation using python-docx
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.ns import qn  # For XML namespace handling
import io
import base64
import re


class DocxGenerator:
    """Advanced DOCX generator with support for complex layouts"""

    def __init__(self):
        self.document = None
        self.styles = None
        self.current_section = None

    def create_document(self, title="Untitled Document", author="Unknown"):
        """Create a new document with basic properties"""
        self.document = Document()
        self.styles = self.document.styles

        # Set document properties
        self.document.core_properties.title = title
        self.document.core_properties.author = author
        self.document.core_properties.comments = "Generated with DocX Generator"

        # Store current section reference
        self.current_section = self.document.sections[0]

        return self.document

    def add_header(self, text, alignment="center", font_size=10):
        """Add header to current section"""
        if not text:
            return

        header = self.current_section.header
        header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        header_para.text = text

        # Set alignment
        alignment_map = {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT
        }
        header_para.alignment = alignment_map.get(alignment.lower(), WD_ALIGN_PARAGRAPH.CENTER)

        # Style header
        run = header_para.runs[0] if header_para.runs else header_para.add_run()
        run.font.size = Pt(font_size)
        run.font.color.rgb = RGBColor(0x42, 0x42, 0x42)
        run.font.name = 'Arial'

    def add_footer(self, text, alignment="center", font_size=10):
        """Add footer to current section"""
        if not text:
            return

        footer = self.current_section.footer
        footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        footer_para.text = text

        # Set alignment
        alignment_map = {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT
        }
        footer_para.alignment = alignment_map.get(alignment.lower(), WD_ALIGN_PARAGRAPH.CENTER)

        # Style footer
        run = footer_para.runs[0] if footer_para.runs else footer_para.add_run()
        run.font.size = Pt(font_size)
        run.font.color.rgb = RGBColor(0x42, 0x42, 0x42)
        run.font.name = 'Arial'

    def add_page_numbers(self, format_type="simple"):
        """Add page numbers to footer"""
        if format_type == "none":
            return

        footer = self.current_section.footer
        page_num_para = footer.add_paragraph()
        page_num_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add page number field
        run = page_num_para.add_run()
        run.font.size = Pt(9)

        if format_type == "simple":
            run.text = "Page "
            self._add_page_number_field(run)
        elif format_type == "total":
            run.text = "Page "
            self._add_page_number_field(run)
            run.add_text(" of ")
            self._add_num_pages_field(run)

    def _add_page_number_field(self, run):
        """Add page number field to run"""
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')

        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = "PAGE"

        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')

        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)

    def _add_num_pages_field(self, run):
        """Add number of pages field to run"""
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')

        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = "NUMPAGES"

        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')

        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)

    def process_html_content(self, html_content, enable_page_breaks=True):
        """Process HTML content and add to document"""
        try:
            from bs4 import BeautifulSoup

            soup = BeautifulSoup(html_content, 'html.parser')

            # Process each element
            for element in soup.find_all():
                self._process_element(element, enable_page_breaks)

        except ImportError:
            # Fallback if BeautifulSoup not available
            self._add_plain_text(html_content)
        except Exception as e:
            print(f"Error processing HTML: {e}")
            self._add_plain_text(html_content)

    def _process_element(self, element, enable_page_breaks):
        """Process individual HTML element"""
        tag = element.name.lower()

        # Skip if element is nested within another processed element
        if element.parent and element.parent.name in ['p', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
            return

        element_handlers = {
            'h1': lambda: self._add_heading(element, 1),
            'h2': lambda: self._add_heading(element, 2),
            'h3': lambda: self._add_heading(element, 3),
            'h4': lambda: self._add_heading(element, 4),
            'h5': lambda: self._add_heading(element, 5),
            'h6': lambda: self._add_heading(element, 6),
            'p': lambda: self._add_paragraph(element),
            'ul': lambda: self._add_list(element, bulleted=True),
            'ol': lambda: self._add_list(element, bulleted=False),
            'li': lambda: None,  # Handled by list processing
            'table': lambda: self._add_table(element),
            'hr': lambda: self._add_horizontal_rule(enable_page_breaks),
            'br': lambda: self._add_line_break(),
            'strong': lambda: self._add_formatted_text(element, bold=True),
            'b': lambda: self._add_formatted_text(element, bold=True),
            'em': lambda: self._add_formatted_text(element, italic=True),
            'i': lambda: self._add_formatted_text(element, italic=True),
            'u': lambda: self._add_formatted_text(element, underline=True),
        }

        handler = element_handlers.get(tag)
        if handler:
            handler()
        elif tag in ['div', 'span', 'section', 'article']:
            # Process container elements
            self._process_container(element, enable_page_breaks)

    def _add_heading(self, element, level):
        """Add heading to document"""
        text = element.get_text().strip()
        if not text:
            return

        style_map = {
            1: 'Heading 1',
            2: 'Heading 2',
            3: 'Heading 3',
            4: 'Heading 4',
            5: 'Heading 5',
            6: 'Heading 6'
        }

        style = style_map.get(level, 'Heading 1')
        para = self.document.add_paragraph(text, style=style)

        # Apply inline formatting if present
        self._apply_inline_formatting(para, element)

    def _add_paragraph(self, element):
        """Add paragraph to document"""
        text = element.get_text().strip()
        if not text:
            return

        para = self.document.add_paragraph(text)
        self._apply_inline_formatting(para, element)

    def _add_list(self, element, bulleted=True):
        """Add list to document"""
        items = element.find_all('li', recursive=False)
        if not items:
            return

        style = 'List Bullet' if bulleted else 'List Number'

        for item in items:
            text = item.get_text().strip()
            if text:
                para = self.document.add_paragraph(text, style=style)
                self._apply_inline_formatting(para, item)

    def _add_table(self, element):
        """Add table to document"""
        try:
            rows = element.find_all('tr')
            if not rows:
                return

            # Calculate number of columns
            num_cols = max(len(row.find_all(['td', 'th'])) for row in rows)

            # Create table
            table = self.document.add_table(rows=len(rows), cols=num_cols)
            table.style = 'Light Grid Accent 1'
            table.autofit = False  # Fixed width

            # Set column widths
            for row in table.rows:
                for cell in row.cells:
                    cell.width = Inches(1.5)

            # Fill table
            for i, row in enumerate(rows):
                cells = row.find_all(['td', 'th'])
                for j, cell in enumerate(cells):
                    if i < len(table.rows) and j < len(table.columns):
                        table_cell = table.cell(i, j)
                        table_cell.text = cell.get_text().strip()

                        # Make header cells bold
                        if cell.name == 'th':
                            for paragraph in table_cell.paragraphs:
                                for run in paragraph.runs:
                                    run.font.bold = True

        except Exception as e:
            print(f"Error processing table: {e}")

    def _add_horizontal_rule(self, enable_page_breaks):
        """Add horizontal rule or page break"""
        if enable_page_breaks:
            self.document.add_page_break()
        else:
            # Add horizontal line
            para = self.document.add_paragraph()
            run = para.add_run("_" * 50)
            run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

    def _add_line_break(self):
        """Add line break"""
        para = self.document.add_paragraph()
        para.add_run().add_break()

    def _add_formatted_text(self, element, bold=False, italic=False, underline=False):
        """Add formatted text (for inline elements)"""
        text = element.get_text().strip()
        if not text:
            return

        # Get the last paragraph and add formatted run
        if self.document.paragraphs:
            para = self.document.paragraphs[-1]
            run = para.add_run(text)

            if bold:
                run.font.bold = True
            if italic:
                run.font.italic = True
            if underline:
                run.font.underline = True

    def _apply_inline_formatting(self, paragraph, element):
        """Apply inline formatting to paragraph based on HTML element"""
        if not paragraph.runs:
            return

        # Check for inline formatting in element
        has_bold = element.find(['strong', 'b'])
        has_italic = element.find(['em', 'i'])
        has_underline = element.find('u')

        for run in paragraph.runs:
            if has_bold:
                run.font.bold = True
            if has_italic:
                run.font.italic = True
            if has_underline:
                run.font.underline = True

    def _process_container(self, element, enable_page_breaks):
        """Process container elements (div, span, etc.)"""
        # Process children
        for child in element.children:
            if hasattr(child, 'name'):
                self._process_element(child, enable_page_breaks)

    def _add_plain_text(self, text):
        """Add plain text as fallback"""
        if text.strip():
            self.document.add_paragraph(text.strip())

    def add_table_of_contents(self):
        """Add table of contents placeholder"""
        toc_heading = self.document.add_paragraph("Table of Contents", style='Heading 1')
        toc_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add placeholder text
        toc_para = self.document.add_paragraph()
        toc_para.add_run("[Table of Contents will be generated here]")
        toc_para.italic = True

        # Add page break
        self.document.add_page_break()

    def add_cover_page(self, title, subtitle=None, author=None, date=None):
        """Add professional cover page"""
        # Add section for cover page
        new_section = self.document.add_section()
        new_section.start_type  # New page

        # Title
        title_para = new_section.add_paragraph()
        title_run = title_para.add_run(title)
        title_run.font.size = Pt(24)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(0x1E, 0x88, 0xE5)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Subtitle
        if subtitle:
            subtitle_para = new_section.add_paragraph()
            subtitle_run = subtitle_para.add_run(subtitle)
            subtitle_run.font.size = Pt(16)
            subtitle_run.font.italic = True
            subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Author
        if author:
            author_para = new_section.add_paragraph()
            author_run = author_para.add_run(f"Author: {author}")
            author_run.font.size = Pt(12)
            author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Date
        if date:
            date_para = new_section.add_paragraph()
            date_run = date_para.add_run(f"Date: {date}")
            date_run.font.size = Pt(12)
            date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add spacing
        for _ in range(3):
            new_section.add_paragraph()

    def save_to_bytes(self):
        """Save document to bytes for download"""
        doc_bytes = io.BytesIO()
        self.document.save(doc_bytes)
        doc_bytes.seek(0)
        return doc_bytes.read()

    def save_to_base64(self):
        """Save document and return as base64 encoded string"""
        doc_bytes = self.save_to_bytes()
        return base64.b64encode(doc_bytes).decode('utf-8')


class MarkdownConverter:
    """Markdown to HTML converter with advanced features"""

    @staticmethod
    def convert_to_html(markdown_text):
        """Convert markdown to HTML"""
        try:
            import markdown2
            extras = [
                "tables",
                "fenced-code-blocks",
                "code-friendly",
                "footnotes",
                "strike",
                "task_list"
            ]
            html = markdown2.markdown(markdown_text, extras=extras)
            return html
        except ImportError:
            # Fallback to simple conversion
            return MarkdownConverter._simple_markdown_to_html(markdown_text)
        except Exception as e:
            print(f"Markdown conversion error: {e}")
            return f"<p>Error converting markdown: {str(e)}</p>"

    @staticmethod
    def _simple_markdown_to_html(markdown_text):
        """Simple markdown to HTML conversion fallback"""
        html = markdown_text

        # Headers
        html = re.sub(r'^###### (.*)$', r'<h6>\1</h6>', html, flags=re.MULTILINE)
        html = re.sub(r'^##### (.*)$', r'<h5>\1</h5>', html, flags=re.MULTILINE)
        html = re.sub(r'^#### (.*)$', r'<h4>\1</h4>', html, flags=re.MULTILINE)
        html = re.sub(r'^### (.*)$', r'<h3>\1</h3>', html, flags=re.MULTILINE)
        html = re.sub(r'^## (.*)$', r'<h2>\1</h2>', html, flags=re.MULTILINE)
        html = re.sub(r'^# (.*)$', r'<h1>\1</h1>', html, flags=re.MULTILINE)

        # Bold and italic
        html = re.sub(r'\*\*(.+?)\*\*', r'<strong>\1</strong>', html)
        html = re.sub(r'\*(.+?)\*', r'<em>\1</em>', html)

        # Links
        html = re.sub(r'\[([^\]]+)\]\(([^\)]+)\)', r'<a href="\2">\1</a>', html)

        # Line breaks to paragraphs
        paragraphs = html.split('\n\n')
        html = '\n'.join(f'<p>{p}</p>' for p in paragraphs if p.strip())

        return html


def generate_docx_from_markdown(markdown_content, options=None):
    """
    Generate DOCX from markdown content with options

    Args:
        markdown_content (str): Markdown content
        options (dict): Generation options
            - title: Document title
            - author: Document author
            - header: Header text
            - footer: Footer text
            - page_number_format: 'none', 'simple', 'total'
            - add_toc: Add table of contents
            - enable_page_breaks: Enable page breaks

    Returns:
        str: Base64 encoded DOCX content
    """
    if options is None:
        options = {}

    # Initialize generator
    generator = DocxGenerator()

    # Create document
    generator.create_document(
        title=options.get('title', 'Untitled Document'),
        author=options.get('author', 'Unknown Author')
    )

    # Add header
    generator.add_header(
        text=options.get('header', ''),
        alignment=options.get('header_alignment', 'center'),
        font_size=options.get('header_font_size', 10)
    )

    # Add footer
    generator.add_footer(
        text=options.get('footer', ''),
        alignment=options.get('footer_alignment', 'center'),
        font_size=options.get('footer_font_size', 10)
    )

    # Add page numbers
    generator.add_page_numbers(
        format_type=options.get('page_number_format', 'none')
    )

    # Convert markdown to HTML
    html_content = MarkdownConverter.convert_to_html(markdown_content)

    # Process content
    generator.process_html_content(
        html_content,
        enable_page_breaks=options.get('enable_page_breaks', True)
    )

    # Add table of contents if requested
    if options.get('add_toc', False):
        generator.add_table_of_contents()

    # Return as base64
    return generator.save_to_base64()


def generate_docx_from_html(html_content, options=None):
    """
    Generate DOCX from HTML content with options

    Args:
        html_content (str): HTML content
        options (dict): Generation options (same as markdown)

    Returns:
        str: Base64 encoded DOCX content
    """
    if options is None:
        options = {}

    # Initialize generator
    generator = DocxGenerator()

    # Create document
    generator.create_document(
        title=options.get('title', 'Untitled Document'),
        author=options.get('author', 'Unknown Author')
    )

    # Add header
    generator.add_header(
        text=options.get('header', ''),
        alignment=options.get('header_alignment', 'center'),
        font_size=options.get('header_font_size', 10)
    )

    # Add footer
    generator.add_footer(
        text=options.get('footer', ''),
        alignment=options.get('footer_alignment', 'center'),
        font_size=options.get('footer_font_size', 10)
    )

    # Add page numbers
    generator.add_page_numbers(
        format_type=options.get('page_number_format', 'none')
    )

    # Process content
    generator.process_html_content(
        html_content,
        enable_page_breaks=options.get('enable_page_breaks', True)
    )

    # Add table of contents if requested
    if options.get('add_toc', False):
        generator.add_table_of_contents()

    # Return as base64
    return generator.save_to_base64()


if __name__ == "__main__":
    # Example usage
    markdown_content = """
# Sample Document

This is a **test document** generated with *python-docx*.

## Features

- Markdown support
- HTML support
- Complex layouts
- Tables and lists

| Feature | Status |
|---------|--------|
| Markdown | ✅ |
| HTML | ✅ |
    """

    options = {
        "title": "Sample Document",
        "author": "DocX Generator",
        "header": "Confidential Document",
        "footer": "Generated with DocX Generator",
        "page_number_format": "simple",
        "add_toc": True,
        "enable_page_breaks": True
    }

    # Generate and save
    docx_base64 = generate_docx_from_markdown(markdown_content, options)

    # Decode and save to file (for testing outside browser)
    import base64
    docx_bytes = base64.b64decode(docx_base64)

    with open("output.docx", "wb") as f:
        f.write(docx_bytes)

    print("Document generated successfully: output.docx")